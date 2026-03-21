"""
PDF转Word工具 - 轻量版
使用PyMuPDF替代pdf2docx，不需要OpenCV
"""

import os
import sys
import re
from typing import Optional, List, Dict, Tuple, Callable
from dataclasses import dataclass

# ==================== 核心依赖 ====================

import fitz  # PyMuPDF
from docx import Document


class PDFConverter:
    """PDF转Word转换器（轻量版）"""

    def __init__(self):
        self.cancelled = False

    def convert(
        self,
        pdf_path: str,
        output_path: str,
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> bool:
        self.cancelled = False

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")

        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        return self._convert_with_pymupdf(pdf_path, output_path, progress_callback)

    def _convert_with_pymupdf(self, pdf_path: str, output_path: str, progress_callback=None) -> bool:
        try:
            doc = Document()
            pdf = fitz.open(pdf_path)
            total_pages = len(pdf)

            for i, page in enumerate(pdf):
                if self.cancelled:
                    pdf.close()
                    return False

                if progress_callback:
                    progress_callback(i + 1, total_pages)

                # 提取文本
                text = page.get_text()
                if text:
                    for para in text.split('\n'):
                        if para.strip():
                            doc.add_paragraph(para.strip())

                # 每页之后添加分页符
                if i < total_pages - 1:
                    doc.add_page_break()

            pdf.close()
            doc.save(output_path)
            return True

        except Exception as e:
            raise RuntimeError(f"PDF转换失败: {str(e)}")

    def cancel(self):
        self.cancelled = True

    @staticmethod
    def get_page_count(pdf_path: str) -> int:
        try:
            pdf = fitz.open(pdf_path)
            count = len(pdf)
            pdf.close()
            return count
        except:
            return 0


# ==================== 搜索替换模块 ====================

from docx.text.paragraph import Paragraph


@dataclass
class MatchResult:
    paragraph_index: int
    text: str
    match_text: str
    start_pos: int
    end_pos: int
    context: str
    location: str


@dataclass
class ReplacementPreview:
    match: MatchResult
    replacement: str
    before: str
    after: str


class SearchReplaceEngine:
    """搜索替换引擎"""

    def __init__(self):
        self.document = None
        self.paragraphs: List[Tuple[int, Paragraph, str]] = []

    def load_document(self, docx_path: str) -> bool:
        try:
            self.document = Document(docx_path)
            self._extract_paragraphs()
            return True
        except Exception as e:
            raise RuntimeError(f"加载文档失败: {str(e)}")

    def _extract_paragraphs(self):
        self.paragraphs = []
        if not self.document:
            return
        for i, para in enumerate(self.document.paragraphs):
            self.paragraphs.append((i, para, "正文"))
        for table_idx, table in enumerate(self.document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        global_idx = -(len(self.paragraphs) + 1)
                        location = f"表格{table_idx + 1}-行{row_idx + 1}-列{cell_idx + 1}"
                        self.paragraphs.append((global_idx, para, location))

    def search(self, keyword: str, case_sensitive: bool = False, whole_word: bool = False) -> List[MatchResult]:
        if not self.document:
            raise RuntimeError("请先加载文档")
        if not keyword:
            return []

        results = []
        flags = 0 if case_sensitive else re.IGNORECASE
        if whole_word:
            pattern = re.compile(r'\b' + re.escape(keyword) + r'\b', flags)
        else:
            pattern = re.compile(re.escape(keyword), flags)

        for para_idx, para, location in self.paragraphs:
            text = para.text
            if not text:
                continue
            for match in pattern.finditer(text):
                start = max(0, match.start() - 20)
                end = min(len(text), match.end() + 20)
                context = text[start:end]
                if start > 0:
                    context = "..." + context
                if end < len(text):
                    context = context + "..."
                result = MatchResult(
                    paragraph_index=para_idx, text=text, match_text=match.group(),
                    start_pos=match.start(), end_pos=match.end(), context=context, location=location
                )
                results.append(result)
        return results

    def preview_replacements(self, keyword: str, replacement: str, case_sensitive: bool = False, whole_word: bool = False) -> List[ReplacementPreview]:
        matches = self.search(keyword, case_sensitive, whole_word)
        previews = []
        for match in matches:
            before = match.text
            after = before[:match.start_pos] + replacement + before[match.end_pos:]
            previews.append(ReplacementPreview(match=match, replacement=replacement, before=before, after=after))
        return previews

    def replace(self, keyword: str, replacement: str, case_sensitive: bool = False, whole_word: bool = False, selected_indices: Optional[List[int]] = None) -> int:
        if not self.document:
            raise RuntimeError("请先加载文档")

        matches = self.search(keyword, case_sensitive, whole_word)
        if not matches:
            return 0

        if selected_indices is not None:
            matches = [m for i, m in enumerate(matches) if i in selected_indices]

        para_matches: Dict[int, List[MatchResult]] = {}
        for match in matches:
            if match.paragraph_index not in para_matches:
                para_matches[match.paragraph_index] = []
            para_matches[match.paragraph_index].append(match)

        count = 0
        for para_idx, para_matches_list in para_matches.items():
            para = None
            for p_idx, p, _ in self.paragraphs:
                if p_idx == para_idx:
                    para = p
                    break
            if para is None:
                continue

            para_matches_list.sort(key=lambda m: m.start_pos, reverse=True)
            text = para.text
            for match in para_matches_list:
                text = text[:match.start_pos] + replacement + text[match.end_pos:]
                count += 1

            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)

        return count

    def save_document(self, output_path: str) -> bool:
        if not self.document:
            raise RuntimeError("没有可保存的文档")
        try:
            self.document.save(output_path)
            return True
        except Exception as e:
            raise RuntimeError(f"保存文档失败: {str(e)}")

    def get_document_statistics(self) -> Dict:
        if not self.document:
            return {}
        return {
            "paragraphs": len(self.document.paragraphs),
            "tables": len(self.document.tables),
            "total_paragraphs": len(self.paragraphs)
        }


# ==================== GUI模块 ====================

from PyQt5.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QFileDialog, QProgressBar,
    QGroupBox, QCheckBox, QTableWidget, QTableWidgetItem,
    QHeaderView, QMessageBox, QSplitter, QStatusBar, QPlainTextEdit
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QColor, QBrush


class ConvertWorker(QThread):
    progress = pyqtSignal(int, int)
    finished = pyqtSignal(bool, str)

    def __init__(self, converter: PDFConverter, pdf_path: str, output_path: str):
        super().__init__()
        self.converter = converter
        self.pdf_path = pdf_path
        self.output_path = output_path

    def run(self):
        try:
            def progress_callback(current, total):
                self.progress.emit(current, total)
            success = self.converter.convert(self.pdf_path, self.output_path, progress_callback)
            self.finished.emit(success, "转换完成" if success else "转换已取消")
        except Exception as e:
            self.finished.emit(False, str(e))


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.pdf_converter = PDFConverter()
        self.search_engine = SearchReplaceEngine()
        self.convert_worker: Optional[ConvertWorker] = None
        self.current_pdf_path = ""
        self.current_docx_path = ""
        self.current_output_path = ""
        self.current_matches: List[MatchResult] = []
        self.current_previews: List[ReplacementPreview] = []
        self.init_ui()

    def init_ui(self):
        self.setWindowTitle("PDF转Word工具 - 敏感词替换")
        self.setMinimumSize(1000, 700)
        self.resize(1200, 800)

        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        main_layout = QVBoxLayout(central_widget)
        main_layout.setSpacing(10)
        main_layout.setContentsMargins(10, 10, 10, 10)

        file_group = self.create_file_selection_group()
        main_layout.addWidget(file_group)

        splitter = QSplitter(Qt.Vertical)
        splitter.addWidget(self.create_search_replace_widget())
        splitter.addWidget(self.create_preview_widget())
        splitter.setSizes([400, 300])
        main_layout.addWidget(splitter, 1)

        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(0)
        self.progress_bar.setFormat("%v/%m 页")
        main_layout.addWidget(self.progress_bar)

        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
        self.status_bar.showMessage("就绪")

    def create_file_selection_group(self) -> QGroupBox:
        group = QGroupBox("文件选择")
        layout = QVBoxLayout(group)

        pdf_layout = QHBoxLayout()
        pdf_label = QLabel("PDF文件:")
        pdf_label.setMinimumWidth(80)
        self.pdf_path_edit = QLineEdit()
        self.pdf_path_edit.setPlaceholderText("选择要转换的PDF文件...")
        self.pdf_path_edit.setReadOnly(True)
        pdf_browse_btn = QPushButton("浏览...")
        pdf_browse_btn.clicked.connect(self.browse_pdf_file)
        pdf_browse_btn.setFixedWidth(80)
        pdf_layout.addWidget(pdf_label)
        pdf_layout.addWidget(self.pdf_path_edit, 1)
        pdf_layout.addWidget(pdf_browse_btn)
        layout.addLayout(pdf_layout)

        output_layout = QHBoxLayout()
        output_label = QLabel("输出目录:")
        output_label.setMinimumWidth(80)
        self.output_path_edit = QLineEdit()
        self.output_path_edit.setPlaceholderText("选择输出目录（默认与PDF同目录）...")
        self.output_path_edit.setReadOnly(True)
        output_browse_btn = QPushButton("浏览...")
        output_browse_btn.clicked.connect(self.browse_output_dir)
        output_browse_btn.setFixedWidth(80)
        output_layout.addWidget(output_label)
        output_layout.addWidget(self.output_path_edit, 1)
        output_layout.addWidget(output_browse_btn)
        layout.addLayout(output_layout)

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        self.convert_btn = QPushButton("转换PDF为Word")
        self.convert_btn.setFixedWidth(150)
        self.convert_btn.clicked.connect(self.convert_pdf)
        btn_layout.addWidget(self.convert_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        return group

    def create_search_replace_widget(self) -> QWidget:
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)

        search_group = QGroupBox("搜索与替换设置")
        search_layout = QVBoxLayout(search_group)

        input_layout = QHBoxLayout()
        search_input_layout = QVBoxLayout()
        search_label = QLabel("搜索词:")
        self.search_edit = QLineEdit()
        self.search_edit.setPlaceholderText("输入要搜索的关键词...")
        self.search_edit.setMinimumWidth(200)
        self.search_edit.returnPressed.connect(self.search_keyword)
        search_input_layout.addWidget(search_label)
        search_input_layout.addWidget(self.search_edit)

        replace_input_layout = QVBoxLayout()
        replace_label = QLabel("替换为:")
        self.replace_edit = QLineEdit()
        self.replace_edit.setPlaceholderText("输入替换后的文本...")
        self.replace_edit.setMinimumWidth(200)
        replace_input_layout.addWidget(replace_label)
        replace_input_layout.addWidget(self.replace_edit)

        input_layout.addLayout(search_input_layout)
        input_layout.addLayout(replace_input_layout)

        options_layout = QHBoxLayout()
        self.case_sensitive_cb = QCheckBox("区分大小写")
        self.whole_word_cb = QCheckBox("全词匹配")
        options_layout.addWidget(self.case_sensitive_cb)
        options_layout.addWidget(self.whole_word_cb)
        options_layout.addStretch()

        self.search_btn = QPushButton("搜索")
        self.search_btn.setFixedWidth(80)
        self.search_btn.clicked.connect(self.search_keyword)
        options_layout.addWidget(self.search_btn)

        self.preview_btn = QPushButton("预览替换")
        self.preview_btn.setFixedWidth(80)
        self.preview_btn.clicked.connect(self.preview_replacements)
        self.preview_btn.setEnabled(False)
        options_layout.addWidget(self.preview_btn)

        input_layout.addLayout(options_layout)
        search_layout.addLayout(input_layout)

        batch_label = QLabel("批量替换列表（每行一个：搜索词=替换词）:")
        search_layout.addWidget(batch_label)

        self.batch_edit = QPlainTextEdit()
        self.batch_edit.setPlaceholderText("例如:\n张三=李四\n电话=联系方式")
        self.batch_edit.setMaximumHeight(80)
        search_layout.addWidget(self.batch_edit)

        layout.addWidget(search_group)

        btn_layout = QHBoxLayout()
        btn_layout.addStretch()
        self.apply_selected_btn = QPushButton("替换选中项")
        self.apply_selected_btn.setFixedWidth(100)
        self.apply_selected_btn.clicked.connect(self.replace_selected)
        self.apply_selected_btn.setEnabled(False)
        btn_layout.addWidget(self.apply_selected_btn)

        self.apply_all_btn = QPushButton("替换全部")
        self.apply_all_btn.setFixedWidth(100)
        self.apply_all_btn.clicked.connect(self.replace_all)
        self.apply_all_btn.setEnabled(False)
        btn_layout.addWidget(self.apply_all_btn)

        self.batch_replace_btn = QPushButton("批量替换")
        self.batch_replace_btn.setFixedWidth(100)
        self.batch_replace_btn.clicked.connect(self.batch_replace)
        self.batch_replace_btn.setEnabled(False)
        btn_layout.addWidget(self.batch_replace_btn)

        self.save_btn = QPushButton("保存文档")
        self.save_btn.setFixedWidth(100)
        self.save_btn.clicked.connect(self.save_document)
        self.save_btn.setEnabled(False)
        btn_layout.addWidget(self.save_btn)
        btn_layout.addStretch()
        layout.addLayout(btn_layout)

        return widget

    def create_preview_widget(self) -> QWidget:
        widget = QWidget()
        layout = QVBoxLayout(widget)
        layout.setContentsMargins(0, 0, 0, 0)

        preview_group = QGroupBox("搜索结果与替换预览")
        preview_layout = QVBoxLayout(preview_group)

        self.result_label = QLabel("共找到 0 处匹配")
        preview_layout.addWidget(self.result_label)

        self.result_table = QTableWidget()
        self.result_table.setColumnCount(6)
        self.result_table.setHorizontalHeaderLabels(["选择", "位置", "匹配文本", "替换为", "上下文", "预览"])
        header = self.result_table.horizontalHeader()
        header.setSectionResizeMode(0, QHeaderView.Fixed)
        header.setSectionResizeMode(1, QHeaderView.Fixed)
        header.setSectionResizeMode(2, QHeaderView.Interactive)
        header.setSectionResizeMode(3, QHeaderView.Interactive)
        header.setSectionResizeMode(4, QHeaderView.Stretch)
        header.setSectionResizeMode(5, QHeaderView.Interactive)
        self.result_table.setColumnWidth(0, 50)
        self.result_table.setColumnWidth(1, 120)
        self.result_table.setColumnWidth(2, 150)
        self.result_table.setColumnWidth(3, 150)
        self.result_table.setColumnWidth(5, 200)
        preview_layout.addWidget(self.result_table)

        select_layout = QHBoxLayout()
        self.select_all_btn = QPushButton("全选")
        self.select_all_btn.setFixedWidth(60)
        self.select_all_btn.clicked.connect(self.select_all_results)
        select_layout.addWidget(self.select_all_btn)

        self.deselect_all_btn = QPushButton("取消全选")
        self.deselect_all_btn.setFixedWidth(70)
        self.deselect_all_btn.clicked.connect(self.deselect_all_results)
        select_layout.addWidget(self.deselect_all_btn)
        select_layout.addStretch()
        preview_layout.addLayout(select_layout)

        layout.addWidget(preview_group)
        return widget

    def browse_pdf_file(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "选择PDF文件", "", "PDF文件 (*.pdf);;所有文件 (*.*)")
        if file_path:
            self.current_pdf_path = file_path
            self.pdf_path_edit.setText(file_path)
            if not self.output_path_edit.text():
                self.output_path_edit.setText(os.path.dirname(file_path))
            self.status_bar.showMessage(f"已选择PDF文件: {os.path.basename(file_path)}")

    def browse_output_dir(self):
        dir_path = QFileDialog.getExistingDirectory(self, "选择输出目录", "")
        if dir_path:
            self.output_path_edit.setText(dir_path)

    def convert_pdf(self):
        if not self.current_pdf_path:
            QMessageBox.warning(self, "警告", "请先选择PDF文件！")
            return
        if not os.path.exists(self.current_pdf_path):
            QMessageBox.warning(self, "警告", "PDF文件不存在！")
            return

        output_dir = self.output_path_edit.text()
        if not output_dir:
            output_dir = os.path.dirname(self.current_pdf_path)

        pdf_name = os.path.splitext(os.path.basename(self.current_pdf_path))[0]
        self.current_docx_path = os.path.join(output_dir, f"{pdf_name}.docx")
        self.current_output_path = self.current_docx_path

        self.convert_btn.setEnabled(False)
        # self.progress_bar.setVisible(True)
        self.progress_bar.setValue(0)
        self.status_bar.showMessage("正在转换PDF...")

        self.convert_worker = ConvertWorker(self.pdf_converter, self.current_pdf_path, self.current_docx_path)
        self.convert_worker.progress.connect(self.on_convert_progress)
        self.convert_worker.finished.connect(self.on_convert_finished)
        self.convert_worker.start()

    def on_convert_progress(self, current: int, total: int):
        if total > 0:
            self.progress_bar.setMaximum(total)
            self.progress_bar.setValue(current)
            self.status_bar.showMessage(f"正在转换: {current}/{total} 页")

    def on_convert_finished(self, success: bool, message: str):
        # self.progress_bar.setVisible(False)
        self.convert_btn.setEnabled(True)

        if success:
            self.status_bar.showMessage(f"转换完成: {self.current_docx_path}")
            try:
                self.search_engine.load_document(self.current_docx_path)
                self.preview_btn.setEnabled(True)
                self.save_btn.setEnabled(True)
                self.batch_replace_btn.setEnabled(True)
                stats = self.search_engine.get_document_statistics()
                self.result_label.setText(f"文档已加载 - 段落数: {stats.get('paragraphs', 0)}, 表格数: {stats.get('tables', 0)}")
                QMessageBox.information(self, "转换成功", f"PDF已成功转换为Word文档！\n\n输出文件: {self.current_docx_path}")
            except Exception as e:
                QMessageBox.warning(self, "警告", f"加载文档失败: {str(e)}")
        else:
            self.status_bar.showMessage(f"转换失败: {message}")
            QMessageBox.warning(self, "转换失败", message)

    def search_keyword(self):
        keyword = self.search_edit.text().strip()
        if not keyword:
            QMessageBox.warning(self, "警告", "请输入搜索关键词！")
            return
        if not self.search_engine.document:
            QMessageBox.warning(self, "警告", "请先转换PDF文件！")
            return

        try:
            self.current_matches = self.search_engine.search(keyword, self.case_sensitive_cb.isChecked(), self.whole_word_cb.isChecked())
            self.update_result_table()
            self.preview_btn.setEnabled(len(self.current_matches) > 0)
            self.apply_all_btn.setEnabled(len(self.current_matches) > 0)
            self.result_label.setText(f"共找到 {len(self.current_matches)} 处匹配")
        except Exception as e:
            QMessageBox.warning(self, "搜索失败", str(e))

    def preview_replacements(self):
        keyword = self.search_edit.text().strip()
        replacement = self.replace_edit.text()
        if not keyword:
            QMessageBox.warning(self, "警告", "请输入搜索关键词！")
            return
        if not self.search_engine.document:
            QMessageBox.warning(self, "警告", "请先转换PDF文件！")
            return

        try:
            self.current_previews = self.search_engine.preview_replacements(keyword, replacement, self.case_sensitive_cb.isChecked(), self.whole_word_cb.isChecked())
            self.update_preview_table()
            self.apply_selected_btn.setEnabled(len(self.current_previews) > 0)
            self.apply_all_btn.setEnabled(len(self.current_previews) > 0)
            self.result_label.setText(f"共 {len(self.current_previews)} 处可替换")
        except Exception as e:
            QMessageBox.warning(self, "预览失败", str(e))

    def update_result_table(self):
        self.result_table.setRowCount(len(self.current_matches))
        for i, match in enumerate(self.current_matches):
            check_item = QTableWidgetItem()
            check_item.setCheckState(Qt.Checked)
            check_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsUserCheckable)
            self.result_table.setItem(i, 0, check_item)
            self.result_table.setItem(i, 1, QTableWidgetItem(match.location))
            self.result_table.setItem(i, 2, QTableWidgetItem(match.match_text))
            self.result_table.setItem(i, 3, QTableWidgetItem(""))
            self.result_table.setItem(i, 4, QTableWidgetItem(match.context))
            self.result_table.setItem(i, 5, QTableWidgetItem(""))

    def update_preview_table(self):
        self.result_table.setRowCount(len(self.current_previews))
        for i, preview in enumerate(self.current_previews):
            check_item = QTableWidgetItem()
            check_item.setCheckState(Qt.Checked)
            check_item.setFlags(Qt.ItemIsEnabled | Qt.ItemIsUserCheckable)
            self.result_table.setItem(i, 0, check_item)
            self.result_table.setItem(i, 1, QTableWidgetItem(preview.match.location))
            match_item = QTableWidgetItem(preview.match.match_text)
            match_item.setBackground(QBrush(QColor(255, 255, 200)))
            self.result_table.setItem(i, 2, match_item)
            replace_item = QTableWidgetItem(preview.replacement)
            replace_item.setBackground(QBrush(QColor(200, 255, 200)))
            self.result_table.setItem(i, 3, replace_item)
            self.result_table.setItem(i, 4, QTableWidgetItem(preview.match.context))
            preview_text = f"...{preview.after[max(0, preview.match.start_pos-20):preview.match.start_pos + len(preview.replacement) + 20]}..."
            self.result_table.setItem(i, 5, QTableWidgetItem(preview_text))

    def select_all_results(self):
        for i in range(self.result_table.rowCount()):
            item = self.result_table.item(i, 0)
            if item:
                item.setCheckState(Qt.Checked)

    def deselect_all_results(self):
        for i in range(self.result_table.rowCount()):
            item = self.result_table.item(i, 0)
            if item:
                item.setCheckState(Qt.Unchecked)

    def get_selected_indices(self) -> List[int]:
        indices = []
        for i in range(self.result_table.rowCount()):
            item = self.result_table.item(i, 0)
            if item and item.checkState() == Qt.Checked:
                indices.append(i)
        return indices

    def replace_selected(self):
        if not self.current_previews:
            return
        selected_indices = self.get_selected_indices()
        if not selected_indices:
            QMessageBox.warning(self, "警告", "请选择要替换的内容！")
            return

        keyword = self.search_edit.text().strip()
        replacement = self.replace_edit.text()
        reply = QMessageBox.question(self, "确认替换", f"确定要替换选中的 {len(selected_indices)} 处内容吗？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.do_replace(keyword, replacement, selected_indices)

    def replace_all(self):
        if not self.current_previews:
            return
        keyword = self.search_edit.text().strip()
        replacement = self.replace_edit.text()
        reply = QMessageBox.question(self, "确认替换", f"确定要替换全部 {len(self.current_previews)} 处内容吗？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            self.do_replace(keyword, replacement, None)

    def do_replace(self, keyword: str, replacement: str, selected_indices: Optional[List[int]]):
        try:
            count = self.search_engine.replace(keyword, replacement, self.case_sensitive_cb.isChecked(), self.whole_word_cb.isChecked(), selected_indices)
            self.status_bar.showMessage(f"已替换 {count} 处")
            self.search_keyword()
            QMessageBox.information(self, "替换完成", f"成功替换了 {count} 处内容！\n\n请记得保存文档。")
        except Exception as e:
            QMessageBox.warning(self, "替换失败", str(e))

    def batch_replace(self):
        batch_text = self.batch_edit.toPlainText().strip()
        if not batch_text:
            QMessageBox.warning(self, "警告", "请输入批量替换列表！")
            return
        if not self.search_engine.document:
            QMessageBox.warning(self, "警告", "请先转换PDF文件！")
            return

        replace_pairs = []
        for line in batch_text.split('\n'):
            line = line.strip()
            if '=' in line:
                parts = line.split('=', 1)
                if len(parts) == 2:
                    search_word = parts[0].strip()
                    replace_word = parts[1].strip()
                    if search_word:
                        replace_pairs.append((search_word, replace_word))

        if not replace_pairs:
            QMessageBox.warning(self, "警告", "未找到有效的替换规则！")
            return

        preview_text = "以下替换将被执行:\n\n"
        total_matches = 0
        for search_word, replace_word in replace_pairs:
            matches = self.search_engine.search(search_word)
            total_matches += len(matches)
            preview_text += f"'{search_word}' -> '{replace_word}': {len(matches)} 处\n"
        preview_text += f"\n总计: {total_matches} 处将被替换"

        reply = QMessageBox.question(self, "确认批量替换", preview_text + "\n\n确定继续吗？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
        if reply == QMessageBox.Yes:
            total_count = 0
            for search_word, replace_word in replace_pairs:
                count = self.search_engine.replace(search_word, replace_word)
                total_count += count
            self.status_bar.showMessage(f"批量替换完成，共替换 {total_count} 处")
            QMessageBox.information(self, "批量替换完成", f"成功替换了 {total_count} 处内容！\n\n请记得保存文档。")
            self.batch_edit.clear()

    def save_document(self):
        if not self.search_engine.document:
            QMessageBox.warning(self, "警告", "没有可保存的文档！")
            return

        default_name = self.current_output_path or "output.docx"
        file_path, _ = QFileDialog.getSaveFileName(self, "保存Word文档", default_name, "Word文档 (*.docx);;所有文件 (*.*)")
        if file_path:
            try:
                self.search_engine.save_document(file_path)
                self.status_bar.showMessage(f"文档已保存: {file_path}")
                QMessageBox.information(self, "保存成功", f"文档已成功保存！\n\n保存位置: {file_path}")
            except Exception as e:
                QMessageBox.warning(self, "保存失败", str(e))

    def closeEvent(self, event):
        if self.convert_worker and self.convert_worker.isRunning():
            reply = QMessageBox.question(self, "确认退出", "PDF转换正在进行中，确定要退出吗？", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)
            if reply == QMessageBox.No:
                event.ignore()
                return
            self.pdf_converter.cancel()
            self.convert_worker.wait()
        event.accept()


def main():
    app = QApplication(sys.argv)
    app.setStyle('Fusion')
    window = MainWindow()
    window.show()
    sys.exit(app.exec_())


if __name__ == "__main__":
    main()