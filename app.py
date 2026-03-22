"""
PDF转Word工具 - 完整版 (CustomTkinter UI)
使用pdf2docx保留PDF排版格式，CustomTkinter现代化界面
"""

import os
import sys
import re
import threading
from typing import Optional, List, Dict, Tuple
from dataclasses import dataclass

# ==================== 核心依赖 ====================

from pdf2docx import Converter
from docx import Document
import customtkinter as ctk
from tkinter import filedialog, messagebox
from tkinter import ttk

# 设置主题
ctk.set_appearance_mode("System")
ctk.set_default_color_theme("blue")


# ==================== 转换参数配置 ====================

class ConversionSettings:
    """PDF转换参数配置"""

    def __init__(self):
        # 图片定位相关参数
        self.float_image_ignorable_gap = 3.0  # 浮动图片识别阈值（降低以识别更多浮动图片）
        self.clip_image_res_ratio = 3.0       # 图片裁剪分辨率比例

        # 向量图形参数
        self.min_svg_gap_dx = 15.0            # 向量图形水平合并间距
        self.min_svg_gap_dy = 2.0             # 向量图形垂直合并间距

        # 页面边距因子
        self.page_margin_factor_top = 0.5     # 页面上边距因子
        self.page_margin_factor_bottom = 0.5  # 页面下边距因子

        # 其他参数
        self.line_overlap_threshold = 0.9     # 行重叠阈值
        self.shape_min_dimension = 2.0        # 忽略小于此值的形状

    def to_dict(self) -> dict:
        """转换为字典供 Converter 使用"""
        return {
            'float_image_ignorable_gap': self.float_image_ignorable_gap,
            'clip_image_res_ratio': self.clip_image_res_ratio,
            'min_svg_gap_dx': self.min_svg_gap_dx,
            'min_svg_gap_dy': self.min_svg_gap_dy,
            'page_margin_factor_top': self.page_margin_factor_top,
            'page_margin_factor_bottom': self.page_margin_factor_bottom,
            'line_overlap_threshold': self.line_overlap_threshold,
            'shape_min_dimension': self.shape_min_dimension,
        }


# ==================== PDF转换模块 ====================

class PDFConverter:
    """PDF转Word转换器（使用pdf2docx保留格式）"""

    def __init__(self, settings: ConversionSettings = None):
        self.cancelled = False
        self.settings = settings or ConversionSettings()

    def convert(self, pdf_path: str, output_path: str, progress_callback=None) -> bool:
        self.cancelled = False

        if not os.path.exists(pdf_path):
            raise FileNotFoundError(f"PDF文件不存在: {pdf_path}")

        output_dir = os.path.dirname(output_path)
        if output_dir and not os.path.exists(output_dir):
            os.makedirs(output_dir)

        return self._convert_with_pdf2docx(pdf_path, output_path, progress_callback)

    def _convert_with_pdf2docx(self, pdf_path: str, output_path: str, progress_callback=None) -> bool:
        try:
            cv = Converter(pdf_path)
            total_pages = len(cv.pages)

            def internal_progress(page_num):
                if progress_callback:
                    progress_callback(page_num, total_pages)
                if self.cancelled:
                    raise InterruptedError("转换已取消")

            # 使用自定义参数进行转换
            cv.convert(output_path, progress=internal_progress, **self.settings.to_dict())
            cv.close()
            return True
        except InterruptedError:
            return False
        except Exception as e:
            raise RuntimeError(f"PDF转换失败: {str(e)}")

    def cancel(self):
        self.cancelled = True


# ==================== 搜索替换模块 ====================

@dataclass
class MatchResult:
    paragraph_index: int
    text: str
    match_text: str
    start_pos: int
    end_pos: int
    context: str
    location: str


@dataclass
class ReplacementPreview:
    match: MatchResult
    replacement: str
    before: str
    after: str


class SearchReplaceEngine:
    """搜索替换引擎"""

    def __init__(self):
        self.document = None
        self.paragraphs: List[Tuple[int, object, str]] = []

    def load_document(self, docx_path: str) -> bool:
        try:
            self.document = Document(docx_path)
            self._extract_paragraphs()
            return True
        except Exception as e:
            raise RuntimeError(f"加载文档失败: {str(e)}")

    def _extract_paragraphs(self):
        self.paragraphs = []
        if not self.document:
            return
        for i, para in enumerate(self.document.paragraphs):
            self.paragraphs.append((i, para, "正文"))
        for table_idx, table in enumerate(self.document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        global_idx = -(len(self.paragraphs) + 1)
                        location = f"表格{table_idx + 1}-行{row_idx + 1}-列{cell_idx + 1}"
                        self.paragraphs.append((global_idx, para, location))

    def search(self, keyword: str, case_sensitive: bool = False, whole_word: bool = False) -> List[MatchResult]:
        if not self.document:
            raise RuntimeError("请先加载文档")
        if not keyword:
            return []

        results = []
        flags = 0 if case_sensitive else re.IGNORECASE
        if whole_word:
            pattern = re.compile(r'\b' + re.escape(keyword) + r'\b', flags)
        else:
            pattern = re.compile(re.escape(keyword), flags)

        for para_idx, para, location in self.paragraphs:
            text = para.text
            if not text:
                continue
            for match in pattern.finditer(text):
                start = max(0, match.start() - 20)
                end = min(len(text), match.end() + 20)
                context = text[start:end]
                if start > 0:
                    context = "..." + context
                if end < len(text):
                    context = context + "..."
                results.append(MatchResult(
                    paragraph_index=para_idx, text=text, match_text=match.group(),
                    start_pos=match.start(), end_pos=match.end(), context=context, location=location
                ))
        return results

    def preview_replacements(self, keyword: str, replacement: str, case_sensitive: bool = False, whole_word: bool = False) -> List[ReplacementPreview]:
        matches = self.search(keyword, case_sensitive, whole_word)
        previews = []
        for match in matches:
            before = match.text
            after = before[:match.start_pos] + replacement + before[match.end_pos:]
            previews.append(ReplacementPreview(match=match, replacement=replacement, before=before, after=after))
        return previews

    def replace(self, keyword: str, replacement: str, case_sensitive: bool = False, whole_word: bool = False, selected_indices: Optional[List[int]] = None) -> int:
        if not self.document:
            raise RuntimeError("请先加载文档")

        matches = self.search(keyword, case_sensitive, whole_word)
        if not matches:
            return 0

        if selected_indices is not None:
            matches = [m for i, m in enumerate(matches) if i in selected_indices]

        para_matches: Dict[int, List[MatchResult]] = {}
        for match in matches:
            if match.paragraph_index not in para_matches:
                para_matches[match.paragraph_index] = []
            para_matches[match.paragraph_index].append(match)

        count = 0
        for para_idx, para_matches_list in para_matches.items():
            para = None
            for p_idx, p, _ in self.paragraphs:
                if p_idx == para_idx:
                    para = p
                    break
            if para is None:
                continue

            para_matches_list.sort(key=lambda m: m.start_pos, reverse=True)
            text = para.text
            for match in para_matches_list:
                text = text[:match.start_pos] + replacement + text[match.end_pos:]
                count += 1

            for run in para.runs:
                run.text = ""
            if para.runs:
                para.runs[0].text = text
            else:
                para.add_run(text)

        return count

    def save_document(self, output_path: str) -> bool:
        if not self.document:
            raise RuntimeError("没有可保存的文档")
        try:
            self.document.save(output_path)
            return True
        except Exception as e:
            raise RuntimeError(f"保存文档失败: {str(e)}")

    def get_document_statistics(self) -> Dict:
        if not self.document:
            return {}
        return {
            "paragraphs": len(self.document.paragraphs),
            "tables": len(self.document.tables),
        }


# ==================== GUI应用 ====================

class PDFToolApp(ctk.CTk):
    def __init__(self):
        super().__init__()

        self.settings = ConversionSettings()
        self.pdf_converter = PDFConverter(self.settings)
        self.search_engine = SearchReplaceEngine()
        self.current_pdf_path = ""
        self.current_docx_path = ""
        self.current_output_path = ""
        self.current_matches: List[MatchResult] = []
        self.current_previews: List[ReplacementPreview] = []
        self.selected_indices: set = set()

        # 窗口设置
        self.title("PDF转Word工具 - 敏感词替换 (完整版)")
        self.geometry("1100x850")
        self.minsize(900, 700)

        # 创建UI
        self.create_widgets()

    def create_widgets(self):
        # 主容器
        self.main_frame = ctk.CTkFrame(self)
        self.main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # === 文件选择区域 ===
        self.file_frame = ctk.CTkFrame(self.main_frame)
        self.file_frame.pack(fill="x", pady=(0, 10))

        # PDF文件选择
        pdf_row = ctk.CTkFrame(self.file_frame, fg_color="transparent")
        pdf_row.pack(fill="x", padx=10, pady=10)

        ctk.CTkLabel(pdf_row, text="PDF文件:", width=80).pack(side="left")
        self.pdf_path_entry = ctk.CTkEntry(pdf_row, placeholder_text="选择要转换的PDF文件...")
        self.pdf_path_entry.pack(side="left", fill="x", expand=True, padx=(10, 10))
        self.pdf_browse_btn = ctk.CTkButton(pdf_row, text="浏览", width=80, command=self.browse_pdf)
        self.pdf_browse_btn.pack(side="left")

        # 输出目录选择
        output_row = ctk.CTkFrame(self.file_frame, fg_color="transparent")
        output_row.pack(fill="x", padx=10, pady=(0, 10))

        ctk.CTkLabel(output_row, text="输出目录:", width=80).pack(side="left")
        self.output_path_entry = ctk.CTkEntry(output_row, placeholder_text="选择输出目录（默认与PDF同目录）...")
        self.output_path_entry.pack(side="left", fill="x", expand=True, padx=(10, 10))
        self.output_browse_btn = ctk.CTkButton(output_row, text="浏览", width=80, command=self.browse_output)
        self.output_browse_btn.pack(side="left")

        # 转换按钮
        convert_row = ctk.CTkFrame(self.file_frame, fg_color="transparent")
        convert_row.pack(fill="x", padx=10, pady=(0, 10))

        self.convert_btn = ctk.CTkButton(convert_row, text="转换PDF为Word", width=150, command=self.convert_pdf)
        self.convert_btn.pack(side="left", padx=(0, 20))

        # 设置按钮
        self.settings_btn = ctk.CTkButton(convert_row, text="高级设置", width=100, command=self.show_settings_dialog)
        self.settings_btn.pack(side="left")

        # 进度条区域
        progress_frame = ctk.CTkFrame(self.main_frame, fg_color="transparent")
        progress_frame.pack(fill="x", pady=(0, 10))

        self.progress_bar = ctk.CTkProgressBar(progress_frame)
        self.progress_bar.pack(fill="x", side="left", expand=True)
        self.progress_bar.set(0)

        self.progress_label = ctk.CTkLabel(progress_frame, text="", width=100)
        self.progress_label.pack(side="right", padx=(10, 0))

        # === 搜索替换区域 ===
        self.search_frame = ctk.CTkFrame(self.main_frame)
        self.search_frame.pack(fill="both", expand=True)

        # 左侧：搜索设置
        left_frame = ctk.CTkFrame(self.search_frame)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 5))

        # 搜索词
        search_row = ctk.CTkFrame(left_frame, fg_color="transparent")
        search_row.pack(fill="x", padx=10, pady=10)

        ctk.CTkLabel(search_row, text="搜索词:").pack(side="left")
        self.search_entry = ctk.CTkEntry(search_row, placeholder_text="输入要搜索的关键词...", width=200)
        self.search_entry.pack(side="left", padx=(10, 5))
        self.search_entry.bind("<Return>", lambda e: self.search_keyword())

        ctk.CTkLabel(search_row, text="替换为:").pack(side="left", padx=(20, 0))
        self.replace_entry = ctk.CTkEntry(search_row, placeholder_text="输入替换后的文本...", width=200)
        self.replace_entry.pack(side="left", padx=(10, 5))

        # 匹配选项
        options_row = ctk.CTkFrame(left_frame, fg_color="transparent")
        options_row.pack(fill="x", padx=10, pady=(0, 10))

        self.case_sensitive_var = ctk.BooleanVar(value=False)
        self.case_sensitive_cb = ctk.CTkCheckBox(options_row, text="区分大小写", variable=self.case_sensitive_var)
        self.case_sensitive_cb.pack(side="left", padx=(0, 20))

        self.whole_word_var = ctk.BooleanVar(value=False)
        self.whole_word_cb = ctk.CTkCheckBox(options_row, text="全词匹配", variable=self.whole_word_var)
        self.whole_word_cb.pack(side="left")

        # 搜索按钮
        btn_row = ctk.CTkFrame(left_frame, fg_color="transparent")
        btn_row.pack(fill="x", padx=10, pady=(0, 10))

        self.search_btn = ctk.CTkButton(btn_row, text="搜索", width=80, command=self.search_keyword)
        self.search_btn.pack(side="left", padx=(0, 10))

        self.preview_btn = ctk.CTkButton(btn_row, text="预览替换", width=80, command=self.preview_replacements, state="disabled")
        self.preview_btn.pack(side="left", padx=(0, 10))

        self.replace_all_btn = ctk.CTkButton(btn_row, text="替换选中项", width=100, command=self.replace_selected, state="disabled")
        self.replace_all_btn.pack(side="left", padx=(0, 10))

        self.replace_all_btn2 = ctk.CTkButton(btn_row, text="替换全部", width=100, command=self.replace_all, state="disabled")
        self.replace_all_btn2.pack(side="left", padx=(0, 10))

        self.save_btn = ctk.CTkButton(btn_row, text="保存文档", width=100, command=self.save_document, state="disabled")
        self.save_btn.pack(side="left")

        # 批量替换
        batch_frame = ctk.CTkFrame(left_frame)
        batch_frame.pack(fill="x", padx=10, pady=(0, 10))

        ctk.CTkLabel(batch_frame, text="批量替换列表（每行一个：搜索词=替换词）:").pack(anchor="w", padx=10, pady=(10, 5))

        self.batch_text = ctk.CTkTextbox(batch_frame, height=60)
        self.batch_text.pack(fill="x", padx=10, pady=(0, 10))
        self.batch_text.insert("1.0", "张三=***\n电话=联系方式\n身份证=证件号码")

        self.batch_replace_btn = ctk.CTkButton(batch_frame, text="批量替换", width=100, command=self.batch_replace, state="disabled")
        self.batch_replace_btn.pack(padx=10, pady=(0, 10))

        # 右侧：结果表格
        right_frame = ctk.CTkFrame(self.search_frame)
        right_frame.pack(side="right", fill="both", expand=True, padx=(5, 0))

        # 结果统计
        self.result_label = ctk.CTkLabel(right_frame, text="共找到 0 处匹配")
        self.result_label.pack(anchor="w", padx=10, pady=10)

        # 表格容器
        table_container = ctk.CTkFrame(right_frame)
        table_container.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        # 创建Treeview表格
        columns = ("select", "location", "match", "replace", "context")
        self.result_tree = ttk.Treeview(table_container, columns=columns, show="headings", height=15)
        
        # 设置列标题
        self.result_tree.heading("select", text="选择")
        self.result_tree.heading("location", text="位置")
        self.result_tree.heading("match", text="匹配文本")
        self.result_tree.heading("replace", text="替换为")
        self.result_tree.heading("context", text="上下文")
        
        # 设置列宽
        self.result_tree.column("select", width=50, anchor="center")
        self.result_tree.column("location", width=120, anchor="center")
        self.result_tree.column("match", width=120, anchor="center")
        self.result_tree.column("replace", width=120, anchor="center")
        self.result_tree.column("context", width=200, anchor="w")
        
        # 添加滚动条
        scrollbar = ttk.Scrollbar(table_container, orient="vertical", command=self.result_tree.yview)
        self.result_tree.configure(yscrollcommand=scrollbar.set)
        
        self.result_tree.pack(side="left", fill="both", expand=True)
        scrollbar.pack(side="right", fill="y")

        # 绑定点击事件
        self.result_tree.bind("<Button-1>", self.on_tree_click)

        # 全选/取消按钮
        select_row = ctk.CTkFrame(right_frame, fg_color="transparent")
        select_row.pack(fill="x", padx=10, pady=(0, 10))

        self.select_all_btn = ctk.CTkButton(select_row, text="全选", width=60, command=self.select_all)
        self.select_all_btn.pack(side="left", padx=(0, 10))

        self.deselect_all_btn = ctk.CTkButton(select_row, text="取消全选", width=70, command=self.deselect_all)
        self.deselect_all_btn.pack(side="left")

        # 状态栏
        self.status_label = ctk.CTkLabel(self, text="就绪", anchor="w")
        self.status_label.pack(fill="x", padx=10, pady=(0, 5))

    def on_tree_click(self, event):
        """处理表格点击事件"""
        region = self.result_tree.identify("region", event.x, event.y)
        if region == "cell":
            column = self.result_tree.identify_column(event.x)
            if column == "#1":  # 选择列
                item = self.result_tree.identify_row(event.y)
                if item:
                    idx = int(item)
                    if idx in self.selected_indices:
                        self.selected_indices.discard(idx)
                        self.result_tree.set(item, "select", "☐")
                    else:
                        self.selected_indices.add(idx)
                        self.result_tree.set(item, "select", "☑")

    def browse_pdf(self):
        file_path = filedialog.askopenfilename(title="选择PDF文件", filetypes=[("PDF文件", "*.pdf"), ("所有文件", "*.*")])
        if file_path:
            self.current_pdf_path = file_path
            self.pdf_path_entry.delete(0, "end")
            self.pdf_path_entry.insert(0, file_path)
            if not self.output_path_entry.get():
                self.output_path_entry.insert(0, os.path.dirname(file_path))
            self.status_label.configure(text=f"已选择: {os.path.basename(file_path)}")

    def browse_output(self):
        dir_path = filedialog.askdirectory(title="选择输出目录")
        if dir_path:
            self.output_path_entry.delete(0, "end")
            self.output_path_entry.insert(0, dir_path)

    def convert_pdf(self):
        if not self.current_pdf_path:
            messagebox.showwarning("警告", "请先选择PDF文件！")
            return

        output_dir = self.output_path_entry.get()
        if not output_dir:
            output_dir = os.path.dirname(self.current_pdf_path)

        pdf_name = os.path.splitext(os.path.basename(self.current_pdf_path))[0]
        self.current_docx_path = os.path.join(output_dir, f"{pdf_name}.docx")
        self.current_output_path = self.current_docx_path

        self.convert_btn.configure(state="disabled")
        self.progress_bar.set(0)
        self.progress_label.configure(text="0/0 页")
        self.status_label.configure(text="正在转换PDF...")

        def do_convert():
            try:
                def progress_callback(current, total):
                    # 使用after确保在主线程更新UI
                    self.after(0, lambda: self.update_progress(current, total))

                success = self.pdf_converter.convert(self.current_pdf_path, self.current_docx_path, progress_callback)
                self.after(0, lambda: self.on_convert_finished(success, "转换完成" if success else "转换已取消"))
            except Exception as e:
                self.after(0, lambda: self.on_convert_finished(False, str(e)))

        threading.Thread(target=do_convert, daemon=True).start()

    def update_progress(self, current: int, total: int):
        """更新进度条"""
        if total > 0:
            progress = current / total
            self.progress_bar.set(progress)
            self.progress_label.configure(text=f"{current}/{total} 页")
            self.status_label.configure(text=f"正在转换: {current}/{total} 页")

    def on_convert_finished(self, success: bool, message: str):
        self.convert_btn.configure(state="normal")

        if success:
            self.progress_bar.set(1.0)
            self.progress_label.configure(text="完成")
            self.status_label.configure(text=f"转换完成: {self.current_docx_path}")
            try:
                self.search_engine.load_document(self.current_docx_path)
                self.preview_btn.configure(state="normal")
                self.save_btn.configure(state="normal")
                self.batch_replace_btn.configure(state="normal")
                stats = self.search_engine.get_document_statistics()
                self.result_label.configure(text=f"文档已加载 - 段落数: {stats.get('paragraphs', 0)}, 表格数: {stats.get('tables', 0)}")
                messagebox.showinfo("转换成功", f"PDF已成功转换为Word文档！\n\n输出文件: {self.current_docx_path}")
            except Exception as e:
                messagebox.showwarning("警告", f"加载文档失败: {str(e)}")
        else:
            self.status_label.configure(text=f"转换失败: {message}")
            messagebox.showwarning("转换失败", message)

    def search_keyword(self):
        keyword = self.search_entry.get().strip()
        if not keyword:
            messagebox.showwarning("警告", "请输入搜索关键词！")
            return

        if not self.search_engine.document:
            messagebox.showwarning("警告", "请先转换PDF文件！")
            return

        self.current_matches = self.search_engine.search(
            keyword,
            case_sensitive=self.case_sensitive_var.get(),
            whole_word=self.whole_word_var.get()
        )

        self.update_result_table()
        self.preview_btn.configure(state="normal" if self.current_matches else "disabled")
        self.result_label.configure(text=f"共找到 {len(self.current_matches)} 处匹配")
        self.status_label.configure(text=f"搜索完成，找到 {len(self.current_matches)} 处匹配")

    def preview_replacements(self):
        keyword = self.search_entry.get().strip()
        replacement = self.replace_entry.get()

        if not keyword:
            messagebox.showwarning("警告", "请输入搜索关键词！")
            return

        if not self.search_engine.document:
            messagebox.showwarning("警告", "请先转换PDF文件！")
            return

        self.current_previews = self.search_engine.preview_replacements(
            keyword, replacement,
            case_sensitive=self.case_sensitive_var.get(),
            whole_word=self.whole_word_var.get()
        )

        self.update_preview_table()
        self.replace_all_btn.configure(state="normal" if self.current_previews else "disabled")
        self.replace_all_btn2.configure(state="normal" if self.current_previews else "disabled")
        self.result_label.configure(text=f"共 {len(self.current_previews)} 处可替换")

    def update_result_table(self):
        """更新搜索结果表格"""
        # 清空表格
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        self.selected_indices.clear()
        
        for i, match in enumerate(self.current_matches):
            self.selected_indices.add(i)
            self.result_tree.insert("", "end", iid=str(i), values=(
                "☑",
                match.location,
                match.match_text,
                "",
                match.context
            ))

    def update_preview_table(self):
        """更新预览表格"""
        # 清空表格
        for item in self.result_tree.get_children():
            self.result_tree.delete(item)
        
        self.selected_indices.clear()
        
        for i, preview in enumerate(self.current_previews):
            self.selected_indices.add(i)
            self.result_tree.insert("", "end", iid=str(i), values=(
                "☑",
                preview.match.location,
                preview.match.match_text,
                preview.replacement,
                preview.match.context
            ))

    def select_all(self):
        """全选"""
        for item in self.result_tree.get_children():
            idx = int(item)
            self.selected_indices.add(idx)
            self.result_tree.set(item, "select", "☑")

    def deselect_all(self):
        """取消全选"""
        self.selected_indices.clear()
        for item in self.result_tree.get_children():
            self.result_tree.set(item, "select", "☐")

    def replace_selected(self):
        """替换选中项"""
        if not self.current_previews:
            return
        
        selected = list(self.selected_indices)
        if not selected:
            messagebox.showwarning("警告", "请选择要替换的内容！")
            return

        keyword = self.search_entry.get().strip()
        replacement = self.replace_entry.get()

        if messagebox.askyesno("确认替换", f"确定要替换选中的 {len(selected)} 处内容吗？"):
            self.do_replace(keyword, replacement, selected)

    def replace_all(self):
        """替换全部"""
        if not self.current_previews:
            return
        
        keyword = self.search_entry.get().strip()
        replacement = self.replace_entry.get()

        if messagebox.askyesno("确认替换", f"确定要替换全部 {len(self.current_previews)} 处内容吗？"):
            self.do_replace(keyword, replacement, None)

    def do_replace(self, keyword: str, replacement: str, selected_indices):
        try:
            count = self.search_engine.replace(
                keyword, replacement,
                case_sensitive=self.case_sensitive_var.get(),
                whole_word=self.whole_word_var.get(),
                selected_indices=selected_indices
            )
            self.status_label.configure(text=f"已替换 {count} 处")
            self.search_keyword()
            messagebox.showinfo("替换完成", f"成功替换了 {count} 处内容！\n\n请记得保存文档。")
        except Exception as e:
            messagebox.showwarning("替换失败", str(e))

    def batch_replace(self):
        batch_text = self.batch_text.get("1.0", "end").strip()
        if not batch_text:
            messagebox.showwarning("警告", "请输入批量替换列表！")
            return

        if not self.search_engine.document:
            messagebox.showwarning("警告", "请先转换PDF文件！")
            return

        replace_pairs = []
        for line in batch_text.split('\n'):
            line = line.strip()
            if '=' in line:
                parts = line.split('=', 1)
                if len(parts) == 2:
                    search_word = parts[0].strip()
                    replace_word = parts[1].strip()
                    if search_word:
                        replace_pairs.append((search_word, replace_word))

        if not replace_pairs:
            messagebox.showwarning("警告", "未找到有效的替换规则！")
            return

        preview_text = "以下替换将被执行:\n\n"
        total_matches = 0
        for search_word, replace_word in replace_pairs:
            matches = self.search_engine.search(search_word)
            total_matches += len(matches)
            preview_text += f"'{search_word}' -> '{replace_word}': {len(matches)} 处\n"
        preview_text += f"\n总计: {total_matches} 处将被替换"

        if messagebox.askyesno("确认批量替换", preview_text + "\n\n确定继续吗？"):
            total_count = 0
            for search_word, replace_word in replace_pairs:
                count = self.search_engine.replace(search_word, replace_word)
                total_count += count
            self.status_label.configure(text=f"批量替换完成，共替换 {total_count} 处")
            messagebox.showinfo("批量替换完成", f"成功替换了 {total_count} 处内容！\n\n请记得保存文档。")

    def save_document(self):
        if not self.search_engine.document:
            messagebox.showwarning("警告", "没有可保存的文档！")
            return

        file_path = filedialog.asksaveasfilename(
            title="保存Word文档",
            defaultextension=".docx",
            initialfile=os.path.basename(self.current_output_path) if self.current_output_path else "output.docx",
            filetypes=[("Word文档", "*.docx"), ("所有文件", "*.*")]
        )

        if file_path:
            try:
                self.search_engine.save_document(file_path)
                self.status_label.configure(text=f"文档已保存: {file_path}")
                messagebox.showinfo("保存成功", f"文档已成功保存！\n\n保存位置: {file_path}")
            except Exception as e:
                messagebox.showwarning("保存失败", str(e))

    def show_settings_dialog(self):
        """显示高级设置对话框"""
        dialog = ctk.CTkToplevel(self)
        dialog.title("转换参数设置")
        dialog.geometry("500x400")
        dialog.transient(self)
        dialog.grab_set()

        # 主容器
        main_frame = ctk.CTkFrame(dialog)
        main_frame.pack(fill="both", expand=True, padx=10, pady=10)

        # 图片定位参数
        image_frame = ctk.CTkFrame(main_frame)
        image_frame.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(image_frame, text="图片定位参数", font=("", 14, "bold")).pack(anchor="w", padx=10, pady=(10, 5))

        # 浮动图片识别阈值
        float_frame = ctk.CTkFrame(image_frame, fg_color="transparent")
        float_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(float_frame, text="浮动图片识别阈值:", width=150).pack(side="left")
        self.float_gap_var = ctk.StringVar(value=str(self.settings.float_image_ignorable_gap))
        float_entry = ctk.CTkEntry(float_frame, textvariable=self.float_gap_var, width=100)
        float_entry.pack(side="left", padx=(10, 0))
        ctk.CTkLabel(float_frame, text="(越小越易识别浮动图片)", text_color="gray").pack(side="left", padx=(10, 0))

        # 图片分辨率比例
        res_frame = ctk.CTkFrame(image_frame, fg_color="transparent")
        res_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(res_frame, text="图片分辨率比例:", width=150).pack(side="left")
        self.res_ratio_var = ctk.StringVar(value=str(self.settings.clip_image_res_ratio))
        res_entry = ctk.CTkEntry(res_frame, textvariable=self.res_ratio_var, width=100)
        res_entry.pack(side="left", padx=(10, 0))
        ctk.CTkLabel(res_frame, text="(影响图片清晰度)", text_color="gray").pack(side="left", padx=(10, 0))

        # 页面边距参数
        margin_frame = ctk.CTkFrame(main_frame)
        margin_frame.pack(fill="x", pady=(0, 10))

        ctk.CTkLabel(margin_frame, text="页面边距参数", font=("", 14, "bold")).pack(anchor="w", padx=10, pady=(10, 5))

        # 上边距因子
        top_frame = ctk.CTkFrame(margin_frame, fg_color="transparent")
        top_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(top_frame, text="上边距因子:", width=150).pack(side="left")
        self.top_margin_var = ctk.StringVar(value=str(self.settings.page_margin_factor_top))
        ctk.CTkEntry(top_frame, textvariable=self.top_margin_var, width=100).pack(side="left", padx=(10, 0))

        # 下边距因子
        bottom_frame = ctk.CTkFrame(margin_frame, fg_color="transparent")
        bottom_frame.pack(fill="x", padx=10, pady=5)
        ctk.CTkLabel(bottom_frame, text="下边距因子:", width=150).pack(side="left")
        self.bottom_margin_var = ctk.StringVar(value=str(self.settings.page_margin_factor_bottom))
        ctk.CTkEntry(bottom_frame, textvariable=self.bottom_margin_var, width=100).pack(side="left", padx=(10, 0))

        # 说明文字
        info_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        info_frame.pack(fill="x", pady=10)
        info_text = """说明:
• 浮动图片识别阈值: 控制图片是否使用绝对定位，值越小识别越敏感
• 图片分辨率比例: 影响转换后图片的质量，值越大越清晰但文件越大
• 页面边距因子: 调整页边距，影响整体布局位置"""
        ctk.CTkLabel(info_frame, text=info_text, justify="left", text_color="gray").pack(anchor="w", padx=10)

        # 按钮
        btn_frame = ctk.CTkFrame(main_frame, fg_color="transparent")
        btn_frame.pack(fill="x", pady=10)

        def apply_settings():
            try:
                self.settings.float_image_ignorable_gap = float(self.float_gap_var.get())
                self.settings.clip_image_res_ratio = float(self.res_ratio_var.get())
                self.settings.page_margin_factor_top = float(self.top_margin_var.get())
                self.settings.page_margin_factor_bottom = float(self.bottom_margin_var.get())
                # 更新转换器设置
                self.pdf_converter.settings = self.settings
                messagebox.showinfo("设置已保存", "参数已更新，将在下次转换时生效。")
                dialog.destroy()
            except ValueError:
                messagebox.showerror("错误", "请输入有效的数字！")

        def reset_settings():
            self.settings = ConversionSettings()
            self.float_gap_var.set(str(self.settings.float_image_ignorable_gap))
            self.res_ratio_var.set(str(self.settings.clip_image_res_ratio))
            self.top_margin_var.set(str(self.settings.page_margin_factor_top))
            self.bottom_margin_var.set(str(self.settings.page_margin_factor_bottom))

        ctk.CTkButton(btn_frame, text="应用", width=80, command=apply_settings).pack(side="left", padx=(10, 10))
        ctk.CTkButton(btn_frame, text="恢复默认", width=80, command=reset_settings).pack(side="left", padx=10)
        ctk.CTkButton(btn_frame, text="取消", width=80, command=dialog.destroy).pack(side="right", padx=10)


def main():
    app = PDFToolApp()
    app.mainloop()


if __name__ == "__main__":
    main()