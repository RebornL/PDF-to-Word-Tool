"""
搜索替换引擎模块
在Word文档中搜索和替换文本，支持预览功能
"""

import re
from typing import List, Dict, Tuple, Optional, Callable
from dataclasses import dataclass
from docx import Document
from docx.text.paragraph import Paragraph
from docx.table import Table, _Cell


@dataclass
class MatchResult:
    """匹配结果"""
    paragraph_index: int      # 段落索引
    text: str                 # 原始文本
    match_text: str           # 匹配到的文本
    start_pos: int            # 开始位置
    end_pos: int              # 结束位置
    context: str              # 上下文（前后各20个字符）
    location: str             # 位置描述（如"正文"或"表格"）


@dataclass
class ReplacementPreview:
    """替换预览"""
    match: MatchResult        # 匹配结果
    replacement: str          # 替换文本
    before: str               # 替换前文本
    after: str                # 替换后文本


class SearchReplaceEngine:
    """搜索替换引擎"""

    def __init__(self):
        self.document: Optional[Document] = None
        self.paragraphs: List[Tuple[int, Paragraph, str]] = []  # (索引, 段落对象, 位置描述)

    def load_document(self, docx_path: str) -> bool:
        """
        加载Word文档

        Args:
            docx_path: Word文档路径

        Returns:
            bool: 是否成功加载
        """
        try:
            self.document = Document(docx_path)
            self._extract_paragraphs()
            return True
        except Exception as e:
            raise RuntimeError(f"加载文档失败: {str(e)}")

    def _extract_paragraphs(self):
        """提取所有段落（包括表格中的段落）"""
        self.paragraphs = []

        if not self.document:
            return

        # 提取正文段落
        for i, para in enumerate(self.document.paragraphs):
            self.paragraphs.append((i, para, "正文"))

        # 提取表格中的段落
        for table_idx, table in enumerate(self.document.tables):
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para_idx, para in enumerate(cell.paragraphs):
                        # 使用负数索引表示表格中的段落
                        global_idx = -(len(self.paragraphs) + 1)
                        location = f"表格{table_idx + 1}-行{row_idx + 1}-列{cell_idx + 1}"
                        self.paragraphs.append((global_idx, para, location))

    def search(
        self,
        keyword: str,
        case_sensitive: bool = False,
        whole_word: bool = False
    ) -> List[MatchResult]:
        """
        在文档中搜索关键词

        Args:
            keyword: 搜索关键词
            case_sensitive: 是否区分大小写
            whole_word: 是否全词匹配

        Returns:
            List[MatchResult]: 匹配结果列表
        """
        if not self.document:
            raise RuntimeError("请先加载文档")

        if not keyword:
            return []

        results = []

        # 构建正则表达式
        flags = 0 if case_sensitive else re.IGNORECASE
        if whole_word:
            pattern = re.compile(r'\b' + re.escape(keyword) + r'\b', flags)
        else:
            pattern = re.compile(re.escape(keyword), flags)

        for para_idx, para, location in self.paragraphs:
            text = para.text
            if not text:
                continue

            for match in pattern.finditer(text):
                # 获取上下文
                start = max(0, match.start() - 20)
                end = min(len(text), match.end() + 20)
                context = text[start:end]

                # 标记匹配位置
                if start > 0:
                    context = "..." + context
                if end < len(text):
                    context = context + "..."

                result = MatchResult(
                    paragraph_index=para_idx,
                    text=text,
                    match_text=match.group(),
                    start_pos=match.start(),
                    end_pos=match.end(),
                    context=context,
                    location=location
                )
                results.append(result)

        return results

    def preview_replacements(
        self,
        keyword: str,
        replacement: str,
        case_sensitive: bool = False,
        whole_word: bool = False
    ) -> List[ReplacementPreview]:
        """
        预览替换结果

        Args:
            keyword: 搜索关键词
            replacement: 替换文本
            case_sensitive: 是否区分大小写
            whole_word: 是否全词匹配

        Returns:
            List[ReplacementPreview]: 替换预览列表
        """
        matches = self.search(keyword, case_sensitive, whole_word)
        previews = []

        for match in matches:
            # 构建替换后的文本
            before = match.text
            after = before[:match.start_pos] + replacement + before[match.end_pos:]

            preview = ReplacementPreview(
                match=match,
                replacement=replacement,
                before=before,
                after=after
            )
            previews.append(preview)

        return previews

    def replace(
        self,
        keyword: str,
        replacement: str,
        case_sensitive: bool = False,
        whole_word: bool = False,
        selected_indices: Optional[List[int]] = None,
        progress_callback: Optional[Callable[[int, int], None]] = None
    ) -> int:
        """
        执行替换操作

        Args:
            keyword: 搜索关键词
            replacement: 替换文本
            case_sensitive: 是否区分大小写
            whole_word: 是否全词匹配
            selected_indices: 选中的匹配索引列表（None表示全部替换）
            progress_callback: 进度回调函数

        Returns:
            int: 替换的数量
        """
        if not self.document:
            raise RuntimeError("请先加载文档")

        matches = self.search(keyword, case_sensitive, whole_word)

        if not matches:
            return 0

        # 如果指定了选中的索引，只替换选中的
        if selected_indices is not None:
            matches = [m for i, m in enumerate(matches) if i in selected_indices]

        # 按段落分组，避免重复处理
        para_matches: Dict[int, List[MatchResult]] = {}
        for match in matches:
            if match.paragraph_index not in para_matches:
                para_matches[match.paragraph_index] = []
            para_matches[match.paragraph_index].append(match)

        total = len(para_matches)
        count = 0

        for idx, (para_idx, para_matches_list) in enumerate(para_matches.items()):
            if progress_callback:
                progress_callback(idx + 1, total)

            # 找到对应的段落对象
            para = None
            for p_idx, p, _ in self.paragraphs:
                if p_idx == para_idx:
                    para = p
                    break

            if para is None:
                continue

            # 从后向前替换，避免位置偏移问题
            para_matches_list.sort(key=lambda m: m.start_pos, reverse=True)

            text = para.text
            for match in para_matches_list:
                text = text[:match.start_pos] + replacement + text[match.end_pos:]
                count += 1

            # 更新段落文本
            self._update_paragraph_text(para, text)

        return count

    def _update_paragraph_text(self, paragraph: Paragraph, new_text: str):
        """
        更新段落文本（保留格式）

        Args:
            paragraph: 段落对象
            new_text: 新文本
        """
        # 清除原有内容
        for run in paragraph.runs:
            run.text = ""

        # 如果有runs，设置到第一个run
        if paragraph.runs:
            paragraph.runs[0].text = new_text
        else:
            # 如果没有runs，直接添加文本
            paragraph.add_run(new_text)

    def save_document(self, output_path: str) -> bool:
        """
        保存文档

        Args:
            output_path: 输出路径

        Returns:
            bool: 是否成功保存
        """
        if not self.document:
            raise RuntimeError("没有可保存的文档")

        try:
            self.document.save(output_path)
            return True
        except Exception as e:
            raise RuntimeError(f"保存文档失败: {str(e)}")

    def get_document_statistics(self) -> Dict:
        """获取文档统计信息"""
        if not self.document:
            return {}

        return {
            "paragraphs": len(self.document.paragraphs),
            "tables": len(self.document.tables),
            "total_paragraphs": len(self.paragraphs)
        }


if __name__ == "__main__":
    # 测试代码
    engine = SearchReplaceEngine()

    # 加载文档
    # engine.load_document("test.docx")

    # 搜索
    # results = engine.search("测试", case_sensitive=False)
    # for r in results:
    #     print(f"找到: {r.match_text} 位置: {r.location}")

    # 预览替换
    # previews = engine.preview_replacements("测试", "替换")
    # for p in previews:
    #     print(f"替换前: {p.before}")
    #     print(f"替换后: {p.after}")

    # 执行替换
    # count = engine.replace("测试", "替换")
    # print(f"替换了 {count} 处")

    # 保存
    # engine.save_document("output.docx")
    pass
